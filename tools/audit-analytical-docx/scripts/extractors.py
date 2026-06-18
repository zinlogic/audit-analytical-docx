import docx
from typing import Dict, Any, List

def extract_summary_data(file_path: str) -> Dict[str, Any]:
    """
    Extrae metadatos y lista de muestras del archivo de Resumen (Tabla).
    """
    doc = docx.Document(file_path)
    
    analyte_info = ""
    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("Analyte:"):
            analyte_info = text
            break

    if len(doc.tables) < 2:
        raise ValueError("El archivo de resumen debe contener al menos 2 tablas.")
        
    metadata_table = doc.tables[0]
    samples_table = doc.tables[1]
    
    # Extraer metadatos
    metadata = {}
    for row in metadata_table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        # Las celdas de metadatos suelen venir en pares clave-valor
        if len(cells) >= 2:
            metadata[cells[0]] = cells[1]
        if len(cells) >= 4:
            metadata[cells[2]] = cells[3]
            
    # Extraer muestras
    samples = []
    headers = [cell.text.strip() for cell in samples_table.rows[0].cells]
    
    for idx, row in enumerate(samples_table.rows[1:]):
        cells = [cell.text.strip() for cell in row.cells]
        if not any(cells): # Saltar filas vacías
            continue
        sample_data = {"id": idx + 1}
        sample_data.update(dict(zip(headers, cells)))
        samples.append(sample_data)
        
    return {
        "metadata": metadata,
        "samples": samples,
        "analyte": analyte_info
    }

def extract_chromatogram_data(file_path: str) -> List[Dict[str, Any]]:
    """
    Extrae los datos de los cromatogramas (Detalle).
    """
    doc = docx.Document(file_path)
    
    # Cada muestra parece tener 4 tablas.
    # En la exploración vimos que la tabla 0 es metadatos de la inyección,
    # la tabla 1 tiene el IS y el Analito Objetivo,
    # la tabla 2 está vacía (posiblemente espacio para el gráfico),
    # y la tabla 3 o 4 podría ser el siguiente.
    # Así que iremos iterando sobre todas las tablas en grupos lógicos o buscando patrones.
    
    samples = []
    
    # Una forma más robusta es buscar la tabla que tiene "Sample Name" y "Batch Name"
    # y luego la tabla siguiente que tiene "Internal Standard" y "Target Analyte"
    
    i = 0
    while i < len(doc.tables):
        table = doc.tables[i]
        
        # Verificar si es una tabla de metadatos de inyección
        if len(table.rows) > 0 and len(table.rows[0].cells) >= 2:
            first_cell_text = table.rows[0].cells[0].text.strip()
            
            if first_cell_text == "Sample Name":
                metadata = {}
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if len(cells) >= 2:
                        metadata[cells[0]] = cells[1]
                    if len(cells) >= 4:
                        metadata[cells[2]] = cells[3]
                
                # La tabla inmediatamente posterior debería ser la de resultados
                i += 1
                if i < len(doc.tables):
                    results_table = doc.tables[i]
                    results = {"Internal Standard": {}, "Target Analyte": {}}
                    
                    if len(results_table.rows) >= 4:
                        # headers = results_table.rows[0].cells (Internal Standard, Area, RT, etc.)
                        # Fila 1 es IS
                        is_cells = [c.text.strip() for c in results_table.rows[1].cells]
                        if len(is_cells) >= 5:
                            results["Internal Standard"] = {
                                "Name": is_cells[0],
                                "Area (cps)": is_cells[1],
                                "RT (min)": is_cells[2],
                                "Calc. Conc. (ng/mL)": is_cells[3],
                                "Calc. Conc.": is_cells[4]
                            }
                        
                        # Fila 3 es Target Analyte
                        target_cells = [c.text.strip() for c in results_table.rows[3].cells]
                        if len(target_cells) >= 5:
                            results["Target Analyte"] = {
                                "Name": target_cells[0],
                                "Area (cps)": target_cells[1],
                                "RT (min)": target_cells[2],
                                "Calc. Conc. (ng/mL)": target_cells[3],
                                "Accuracy (%)": target_cells[4]
                            }
                            
                    samples.append({
                        "id": len(samples) + 1,
                        "metadata": metadata,
                        "results": results
                    })
        i += 1

    return samples

if __name__ == "__main__":
    # Test script locally
    import json
    sum_data = extract_summary_data("datos/Tabla Secuencia 1-RE-EZPC.docx")
    det_data = extract_chromatogram_data("datos/Cromatogramas Secuencia 1-RE-EZPC.docx")
    print("Summary Samples:", len(sum_data["samples"]))
    print("Detail Samples:", len(det_data))
